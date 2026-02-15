from pathlib import Path

from docx import Document

from app.ingestion.processors.base import BaseProcessor, TextBlock


class DocxProcessor(BaseProcessor):
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, file_path: Path) -> list[TextBlock]:
        doc = Document(str(file_path))
        blocks: list[TextBlock] = []
        current_heading: str = ""
        current_content: list[str] = []

        for para in doc.paragraphs:
            if para.style and para.style.name.startswith("Heading"):
                # Save previous block
                if current_content:
                    text = "\n".join(current_content).strip()
                    if text:
                        blocks.append(TextBlock(
                            content=text,
                            metadata={
                                "file_type": "docx",
                                "heading": current_heading,
                                "style": para.style.name if para.style else "",
                            },
                        ))
                    current_content = []
                current_heading = para.text.strip()
                current_content.append(para.text)
            else:
                if para.text.strip():
                    current_content.append(para.text)

        # Save last block
        if current_content:
            text = "\n".join(current_content).strip()
            if text:
                blocks.append(TextBlock(
                    content=text,
                    metadata={
                        "file_type": "docx",
                        "heading": current_heading,
                    },
                ))

        # Also extract tables
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                blocks.append(TextBlock(
                    content=table_text,
                    metadata={
                        "file_type": "docx",
                        "content_type": "table",
                        "table_index": i,
                    },
                ))

        if not blocks:
            blocks.append(TextBlock(
                content="[No extractable text found in DOCX]",
                metadata={"file_type": "docx"},
            ))

        return blocks
